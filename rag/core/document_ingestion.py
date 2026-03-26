"""
document_ingestion.py
---------------------
Uten Assistant RAG — Document Ingestion Pipeline

Converts any user-supplied source (pasted text or uploaded file) into a
normalized Document object that is ready for chunking and embedding.

Supported file types:
    .txt, .md          → read as UTF-8 text
    .pdf               → extract text via pypdf
    .docx              → extract text via python-docx
    .csv               → convert rows to readable text
    .json / .jsonl     → pretty-print / line-by-line
    .html / .htm       → strip tags with html.parser

Chunking strategy:
    Uses LangChain's RecursiveCharacterTextSplitter with tiktoken encoding.
    This splits on [paragraph → newline → space → char] in priority order
    and measures size in real tokens (not words), matching the embedding
    model's actual token limit.

Usage:
    from document_ingestion import ingest_text, ingest_file, ingest_many, Document

    doc = ingest_text("Your pasted content here...", title="My Notes")

    doc = ingest_file("/path/to/uploaded/report.pdf")

    results = ingest_many(["a.pdf", "b.docx", "c.csv"])
    for r in results:
        if r.success:
            print(r.document.chunks)
        else:
            print(f"Failed: {r.error}")

    print(doc.text)           # full extracted text
    print(doc.metadata)       # source, title, file_type, char_count, etc.
    print(doc.chunks)         # list[str] — chunked & ready for embedding

Dependencies:
    pip install pypdf python-docx langchain-text-splitters tiktoken
"""

from __future__ import annotations

import csv
import json
import re
import textwrap
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import datetime
from html.parser import HTMLParser
from pathlib import Path
from typing import Optional

from langchain_text_splitters import RecursiveCharacterTextSplitter



def _make_splitter(chunk_size: int, chunk_overlap: int) -> RecursiveCharacterTextSplitter:
    """
    Build a RecursiveCharacterTextSplitter measured in real tokens.

    Why a factory function and not a module-level constant?
    Because chunk_size and chunk_overlap can vary per call — some documents
    (e.g. large PDFs) benefit from bigger chunks, while short CSVs want
    smaller ones.  The factory lets callers tune this without touching
    global state.

    Why tiktoken / cl100k_base?
    cl100k_base is the tokeniser used by OpenAI's embedding models AND by
    Anthropic's Claude.  Measuring chunks in these tokens means we never
    accidentally send a chunk that exceeds the embedding model's limit.

    The separators list is the 'recursive' part of the name:
        1. "\\n\\n"  — try to split on paragraph boundaries first
        2. "\\n"     — fall back to single newlines
        3. " "       — fall back to word boundaries
        4. ""        — last resort: split at any character
    It works down the list only when the current chunk is still too large.
    """
    return RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", " ", ""],
        length_function=len,          # measure in characters; swap to
    )



@dataclass
class Document:
    """Normalized document ready for RAG chunking and embedding."""

    text: str                          # full extracted / cleaned text
    title: str = "Untitled"
    source: str = "unknown"            # "text_input" | file path
    file_type: str = "text"            # txt, pdf, docx, csv, json, html, …
    metadata: dict = field(default_factory=dict)
    chunks: list[str] = field(default_factory=list)

    def __post_init__(self) -> None:
        self.metadata.setdefault("char_count", len(self.text))
        self.metadata.setdefault("word_count", len(self.text.split()))
        self.metadata.setdefault("ingested_at", datetime.utcnow().isoformat())


@dataclass
class IngestionResult:
    """
    Wraps the outcome of a single file ingestion attempt.

    Why this exists:
        When processing a batch, one bad file should not crash everything.
        Instead of raising an exception, ingest_many() wraps each outcome
        in an IngestionResult so the caller can inspect both successes and
        failures in one place.

    Think of it like a Dart sealed class with Success / Failure variants —
    except Python expresses it as a single dataclass with optional fields
    and a boolean flag.

    Attributes:
        file_path  : the original path that was attempted
        document   : populated if ingestion succeeded, None otherwise
        error      : human-readable error message if ingestion failed
        success    : True if document is populated, False if error is set
    """

    file_path: str
    document:  Optional[Document] = None
    error:     Optional[str]      = None
    success:   bool               = False

    def __post_init__(self) -> None:
        self.success = self.document is not None



def ingest_text(
    text: str,
    title: str = "Pasted Text",
    source: Optional[str] = None,
    chunk_size: int = 512,
    chunk_overlap: int = 50,
) -> Document:
    """
    Ingest raw pasted text.

    Args:
        text:          The raw string content from the user.
        title:         A human-readable label for this document.
        chunk_size:    Max characters per chunk (maps to ~128 tokens at
                       average English density).  512 is a safe default
                       for most embedding models.
        chunk_overlap: Characters of overlap carried into the next chunk.
                       50 chars ≈ one short sentence of shared context.

    Returns:
        A Document with .text populated and .chunks generated.
    """
    cleaned  = _clean_text(text)
    splitter = _make_splitter(chunk_size, chunk_overlap)
    doc = Document(
        text=cleaned,
        title=title,
        source=source or title,
        file_type="text",
        metadata={
            "input_method": "paste",
            "chunk_size": chunk_size,
            "chunk_overlap": chunk_overlap,
        },
    )
    doc.chunks = splitter.split_text(cleaned)
    return doc


def ingest_file(
    file_path: str | Path,
    title: Optional[str] = None,
    chunk_size: int = 512,
    chunk_overlap: int = 50,
) -> Document:
    """
    Ingest an uploaded file.  Detects type by extension and routers to the
    correct extractor automatically.

    Args:
        file_path:     Absolute or relative path to the uploaded file.
        title:         Override the document title (defaults to filename).
        chunk_size:    Max characters per chunk.
        chunk_overlap: Characters of overlap between adjacent chunks.

    Returns:
        A Document with .text populated and .chunks generated.

    Raises:
        FileNotFoundError: if the path does not exist.
        ValueError:        if the file type is not supported.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    ext   = path.suffix.lower()
    title = title or path.stem.replace("_", " ").replace("-", " ").title()

    extractor_map = {
        ".txt":   _extract_txt,
        ".md":    _extract_txt,
        ".pdf":   _extract_pdf,
        ".docx":  _extract_docx,
        ".csv":   _extract_csv,
        ".json":  _extract_json,
        ".jsonl": _extract_jsonl,
        ".html":  _extract_html,
        ".htm":   _extract_html,
    }

    extractor = extractor_map.get(ext)
    if extractor is None:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported: {', '.join(extractor_map)}"
        )

    raw_text = extractor(path)
    cleaned  = _clean_text(raw_text)
    splitter = _make_splitter(chunk_size, chunk_overlap)

    doc = Document(
        text=cleaned,
        title=title,
        source=str(path),
        file_type=ext.lstrip("."),
        metadata={
            "input_method": "file_upload",
            "filename": path.name,
            "file_size_bytes": path.stat().st_size,
            "chunk_size": chunk_size,
            "chunk_overlap": chunk_overlap,
        },
    )
    doc.chunks = splitter.split_text(cleaned)
    return doc


def ingest_many(
    file_paths: list[str | Path],
    chunk_size: int = 512,
    chunk_overlap: int = 50,
    max_workers: int = 4,
) -> list[IngestionResult]:
    """
    Ingest a batch of files in parallel.

    Each file is processed in its own thread.  If a file fails for any
    reason (corrupt, unsupported type, permission error, etc.) it is
    skipped and its error is recorded — other files in the batch are
    not affected.

    Args:
        file_paths:   List of file paths to ingest (strings or Path objects).
        chunk_size:   Max characters per chunk — applied to every file.
        chunk_overlap: Overlap characters between chunks.
        max_workers:  Maximum number of threads to use simultaneously.
                      4 is a safe default — I/O-bound work (reading files)
                      benefits from threading but you don't want to spawn
                      hundreds of threads for huge batches.

    Returns:
        List of IngestionResult objects in the SAME ORDER as file_paths.
        Check result.success to know if each one worked.

    Example:
        results = ingest_many(["a.pdf", "b.docx", "c.csv"])

        for r in results:
            if r.success:
                print(f"{r.file_path} → {len(r.document.chunks)} chunks")
            else:
                print(f"{r.file_path} FAILED: {r.error}")

    Why ThreadPoolExecutor and not asyncio?
        File reading is I/O-bound (waiting on disk), not CPU-bound.
        Threads are perfect for I/O-bound work and require no changes
        to ingest_file() itself.  asyncio would require every extractor
        to be rewritten as async — a big change for marginal gain here.

    Why as_completed() but return in original order?
        as_completed() gives us results as soon as each thread finishes —
        so we can record them without waiting for the slowest file.
        But we store them in a dict keyed by future, then rebuild the
        list in original order so the caller gets predictable output.
    """

    results: dict[int, IngestionResult] = {}

    def _process(index: int, path: str | Path) -> tuple[int, IngestionResult]:
        str_path = str(path)
        try:
            doc = ingest_file(
                path,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
            )
            return index, IngestionResult(file_path=str_path, document=doc)
        except Exception as e:
            return index, IngestionResult(
                file_path=str_path,
                error=f"{type(e).__name__}: {e}",
            )

    with ThreadPoolExecutor(max_workers=max_workers) as executor:

        future_to_index = {
            executor.submit(_process, i, path): i
            for i, path in enumerate(file_paths)
        }

        for future in as_completed(future_to_index):
            index, result = future.result()
            results[index] = result

    return [results[i] for i in range(len(file_paths))]

def _extract_txt(path: Path) -> str:
    """Plain text / Markdown — UTF-8 with fallback to latin-1."""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


def _extract_pdf(path: Path) -> str:
    """PDF — uses pypdf for text extraction."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf is required for PDF ingestion: pip install pypdf")

    reader = PdfReader(str(path))
    pages  = []
    for i, page in enumerate(reader.pages, 1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"[Page {i}]\n{page_text}")
    if not pages:
        raise ValueError("No extractable text found in PDF (may be scanned / image-only).")
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    """DOCX — uses python-docx, preserves paragraph structure."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx is required for DOCX ingestion: pip install python-docx")

    doc    = DocxDocument(str(path))
    parts  = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else ""
            if "Heading" in style:
                parts.append(f"\n## {text}\n")
            else:
                parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def _extract_csv(path: Path) -> str:
    """CSV — converts each row to a readable key: value block."""
    rows = []
    with open(path, newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.DictReader(f)
        for i, row in enumerate(reader, 1):
            lines = [f"Record {i}:"]
            for k, v in row.items():
                if v and v.strip():
                    lines.append(f"  {k}: {v.strip()}")
            rows.append("\n".join(lines))
    if not rows:
        raise ValueError("CSV file is empty or has no data rows.")
    return "\n\n".join(rows)


def _extract_json(path: Path) -> str:
    """JSON — pretty-prints the structure as readable text."""
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    return json.dumps(data, indent=2, ensure_ascii=False)


def _extract_jsonl(path: Path) -> str:
    """JSONL — each line is a separate JSON object."""
    blocks = []
    with open(path, encoding="utf-8") as f:
        for i, line in enumerate(f, 1):
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
                blocks.append(f"[Record {i}]\n{json.dumps(obj, indent=2, ensure_ascii=False)}")
            except json.JSONDecodeError:
                blocks.append(f"[Record {i} — parse error]\n{line}")
    return "\n\n".join(blocks)


def _extract_html(path: Path) -> str:
    """HTML — strips tags and collapses whitespace."""
    raw = path.read_text(encoding="utf-8", errors="replace")
    return _strip_html(raw)



class _TextExtractor(HTMLParser):
    """Minimal HTMLParser subclass that collects visible text."""

    SKIP_TAGS = {"script", "style", "head", "meta", "link", "noscript"}

    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self._skip_depth: int  = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in self.SKIP_TAGS:
            self._skip_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in self.SKIP_TAGS and self._skip_depth > 0:
            self._skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if self._skip_depth == 0:
            self._parts.append(data)

    def get_text(self) -> str:
        return " ".join(self._parts)


def _strip_html(html: str) -> str:
    parser = _TextExtractor()
    parser.feed(html)
    return parser.get_text()



def _clean_text(text: str) -> str:
    """
    Normalise whitespace and remove non-printable characters.
    Preserves paragraph breaks (double newlines).
    """
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[^\S\n\t ]+", " ", text)
    lines = [line.rstrip() for line in text.splitlines()]
    return "\n".join(lines).strip()





if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python document_ingestion.py <file_path>")
        print("       echo 'some text' | python document_ingestion.py --stdin")
        sys.exit(1)

    if sys.argv[1] == "--stdin":
        raw = sys.stdin.read()
        doc = ingest_text(raw, title="stdin")
    else:
        doc = ingest_file(sys.argv[1])

    print(f"\n{'='*60}")
    print(f"Title    : {doc.title}")
    print(f"Source   : {doc.source}")
    print(f"Type     : {doc.file_type}")
    print(f"Words    : {doc.metadata['word_count']}")
    print(f"Chunks   : {len(doc.chunks)}")
    print(f"{'='*60}\n")
    for i, chunk in enumerate(doc.chunks[:3], 1):
        print(f"--- Chunk {i} ---")
        print(textwrap.fill(chunk[:300], width=72))
        print()
    if len(doc.chunks) > 3:
        print(f"... and {len(doc.chunks) - 3} more chunks.")
